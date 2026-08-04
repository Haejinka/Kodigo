from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Kodigo v0.1.0")
OUT = ROOT / "deliverables" / "Kodigo_Governance_and_Full_Manual.docx"
LOGO = ROOT / "kodigo-ui" / "public" / "kodigo-icon.png"

BLUE = "2E74B5"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "E9F6EE"
AMBER = "FFF4D6"
RED = "FDECEC"
WHITE = "FFFFFF"
BLACK = "1F2937"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in, indent_dxa=120):
    total_dxa = sum(round(w * 1440) for w in widths_in)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_in)):
            dxa = round(width * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width)
            set_cell_margins(cell)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(table):
    if table.rows:
        repeat_header(table.rows[0])


def style_run(run, size=None, bold=None, color=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        style_run(r, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        style_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        style_run(r, italic=italic)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def new_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    new_num_id = max(existing or [0]) + 1
    # Reuse the abstract numbering definition behind Word's built-in List Number style.
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    abstract_id = None
    for num in numbering.findall(qn("w:num")):
        if int(num.get(qn("w:numId"))) == int(style_num_id):
            abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
            break
    if abstract_id is None:
        abstract_id = "0"
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def add_steps(doc, items):
    num_id = new_numbering_instance(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = num_pr.get_or_add_ilvl()
        ilvl.val = 0
        num = num_pr.get_or_add_numId()
        num.val = num_id
        p.add_run(item)


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    style_run(r, bold=True, color=NAVY)
    style_run(p.add_run(text), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths, font_size=8.5, header_fill=LIGHT_BLUE, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        style_run(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if alignments:
                p.alignment = alignments[i]
            r = p.add_run(str(value))
            style_run(r, size=font_size, color=BLACK)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, widths)
    set_repeat_table_header(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_flow_diagram(doc, steps):
    table = doc.add_table(rows=0, cols=1)
    table.style = "Table Grid"
    for idx, (title, desc) in enumerate(steps, start=1):
        row = table.add_row()
        cell = row.cells[0]
        set_cell_shading(cell, LIGHT_BLUE if idx % 2 else PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(f"{idx}. {title}"), size=10.5, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        style_run(p2.add_run(desc), size=8.5, color=BLACK)
        prevent_row_split(row)
        if idx != len(steps):
            arrow = table.add_row().cells[0]
            arrow_p = arrow.paragraphs[0]
            arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_p.paragraph_format.space_after = Pt(0)
            style_run(arrow_p.add_run("↓"), size=12, bold=True, color=BLUE)
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def new_section(doc, title, subtitle=None):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper())
    style_run(r, size=10, bold=True, color=BLUE)
    h = doc.add_paragraph()
    h.style = doc.styles["Heading 1"]
    h.paragraph_format.space_before = Pt(0)
    r = h.add_run(title)
    style_run(r, size=22, bold=True, color=NAVY)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(14)
        style_run(p2.add_run(subtitle), size=11, color=MID_GRAY, italic=True)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (11.5, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.18
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.7)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.2)
    styles["List Number"].paragraph_format.left_indent = Inches(0.375)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)

    for section in doc.sections:
        section.different_first_page_header_footer = True
    return doc


def add_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    style_run(hp.add_run("KODIGO  |  GOVERNANCE, POLICY & USER MANUAL"), size=8, bold=True, color=MID_GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(fp.add_run("Kodigo Capstone Documentation  •  "), size=8, color=MID_GRAY)
    add_field(fp, "PAGE")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    core = doc.core_properties
    core.title = "Kodigo Governance Artifacts and Full User Manual"
    core.subject = "Business flow, COBIT and ITIL artifacts, user manual, policies, data classification, and privacy notice"
    core.author = "Kodigo Capstone Team"
    core.keywords = "Kodigo, POS, inventory, COBIT, ITIL, AUP, BYOD, privacy"

    # Cover: editorial_cover pattern, adapted for a formal manual.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(1.1))
        logo_shape._inline.docPr.set("descr", "Kodigo application logo")
        logo_shape._inline.docPr.set("title", "Kodigo")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("KODIGO"), size=32, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("Governance Artifacts & Full User Manual"), size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Cloud-first point-of-sale, inventory, supplier, and analytics system"), size=12, italic=True, color=MID_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_callout(doc, "Document scope", "Proposed business flow; COBIT-aligned permissions and accountability; ITIL-aligned escalation; full operating manual; AUP and BYOD policies; data classification; and a user-facing privacy notice.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    style_run(p.add_run("Version 1.0  |  2 August 2026"), size=10, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Prepared for capstone evaluation and operational use"), size=9.5, color=MID_GRAY)

    doc.add_page_break()
    add_heading(doc, "Document Control", 1)
    add_table(doc,
              ["Field", "Value"],
              [
                  ("System", "Kodigo multi-store POS and inventory platform"),
                  ("Audience", "Store administrators, cashiers, inventory staff, super administrators, and technical support"),
                  ("Owner", "Kodigo System Owner / Store Administration"),
                  ("Approval", "Capstone adviser and designated system owner"),
                  ("Review cycle", "At least annually and after material role, privacy, or workflow changes"),
                  ("Source basis", "Current React/Vite UI, Supabase migration chain through migration 28, and production operations runbook"),
              ], [1.5, 5.0], font_size=9.5)
    add_heading(doc, "How to Use This Document", 2)
    add_body(doc, "Sections 1–3 are governance and service-management artifacts. Section 4 is the day-to-day manual. Sections 5–7 are policy and privacy artifacts that should be approved and adopted by the deploying organization before production use.")
    add_callout(doc, "Implementation note", "Kodigo is store-scoped. A user sees only stores assigned through membership mappings, and the selected active store controls operational context. The “All Stores” view is for aggregation and cannot process a sale.", fill=AMBER)
    add_heading(doc, "Contents", 1)
    contents = [
        "1. Proposed Business Flow",
        "2. COBIT Artifact — Permission Matrix & Roles and Responsibilities",
        "3. ITIL Artifact — Escalation Path",
        "4. Full Manual Guide",
        "5. Acceptable Use and BYOD Policy",
        "6. Data Classification Table",
        "7. Privacy Notice",
        "Appendix A. Control and Evidence Notes",
    ]
    add_bullets(doc, contents)

    # 1 Business flow
    new_section(doc, "1. Proposed Business Flow", "From access provisioning to sale, stock control, reporting, and continual improvement")
    add_body(doc, "The proposed operating model starts with controlled onboarding and ends with review and corrective action. Each step protects the reliability of sales and inventory records while keeping frontline work fast enough for a busy small retailer.")
    flow_steps = [
        ("Provision and assign access", "An Admin creates cashier or inventory accounts and assigns the correct store; a Super Admin governs invite codes. This establishes least privilege and tenancy before any transaction occurs."),
        ("Authenticate and select store", "The user signs in through Supabase Auth. Kodigo resolves the role and store memberships, then loads a specific store or an authorized aggregate view. Correct context prevents cross-branch posting."),
        ("Prepare the catalog and stock", "Admin or Inventory staff create products, SKUs/barcodes, prices, units, thresholds, categories, and stock adjustments. Accurate master data is required for correct checkout, replenishment, and reporting."),
        ("Process the sale", "Admin or Cashier scans/searches products, confirms quantities, accepts payment, and completes checkout. The transaction service records the sale, items, payment and receipt while deducting stock atomically."),
        ("Handle connectivity and exceptions", "If the network fails, eligible sales or mutations are queued in IndexedDB and replayed on reconnection. Business/RLS errors are shown instead of being misclassified as offline work."),
        ("Monitor stock and replenish", "Kodigo creates low/critical/out-of-stock alerts and recommends restock quantities. Admin reviews suppliers, creates purchase orders, and marks deliveries received so stock and supplier performance stay current."),
        ("Review reports and controls", "Admin reviews dashboard, analytics, rankings, sales and stock-movement reports; Inventory can access inventory-focused reports. Exceptions such as voids, refunds, closeouts, and errors are traceable in logs."),
        ("Improve and govern", "The System Owner reviews incidents, permissions, audit evidence, backup tests, and user feedback. Approved changes move through staging, testing, production deployment, and documented rollback procedures."),
    ]
    add_flow_diagram(doc, flow_steps)
    add_heading(doc, "Why the Flow Matters", 2)
    add_table(doc, ["Control objective", "How the flow supports it"], [
        ("Accuracy", "Catalog setup precedes sales; checkout records item, payment, tax/discount snapshots, receipt, and stock effects as one controlled operation."),
        ("Accountability", "Named users, roles, store membership, sale events, audit logs, and stock adjustment records show who did what and where."),
        ("Continuity", "Offline queues keep authorized operations available during network interruption, while backups and rollback protect recovery."),
        ("Decision support", "Current stock, supplier scoring, alerts, and reports convert transaction data into replenishment and performance decisions."),
    ], [1.6, 4.9], font_size=9.2)

    # 2 COBIT
    new_section(doc, "2. COBIT Artifact — Permission Matrix & Roles and Responsibilities", "Least privilege, segregation of duties, ownership, approval, and accountability")
    add_callout(doc, "Permission legend", "Full Control = administer and configure; Edit = create/change operational records; Approve = authorize sensitive action; Read = view only; Execute = perform a narrow transaction; None = no access. All application access remains limited to assigned stores.")
    add_heading(doc, "2.1 Functional Permission Matrix", 2)
    perm_rows = [
        ("POS checkout", "Execute", "Execute", "None", "None"),
        ("Inventory products & stock", "Full Control", "Read at POS", "Edit", "None"),
        ("Stock adjustments", "Approve/Edit", "None", "Edit", "None"),
        ("Suppliers & purchase orders", "Full Control", "None", "Read only where exposed", "None"),
        ("Dashboard & analytics", "Full Control", "None", "None", "None"),
        ("Reports/export", "Full sales + inventory", "None", "Inventory/stock only", "None"),
        ("Store settings & branding", "Full Control", "None", "None", "None"),
        ("Manage store users", "Full Control in shared stores", "Own account only", "Own account only", "None"),
        ("Invite-code governance", "None", "None", "None", "Full Control"),
        ("Notifications", "Store/admin notices", "Operational notices only", "Limited", "System notices"),
        ("Account password", "Own account", "Own account", "Own account", "Own account"),
        ("Audit/error logs", "Read scoped operational evidence", "None", "None", "System governance view as configured"),
    ]
    add_table(doc, ["Capability", "Admin", "Cashier", "Inventory", "Super Admin"], perm_rows,
              [2.1, 1.1, 1.1, 1.1, 1.1], font_size=7.8,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    add_body(doc, "Important: Super Admin is a platform-governance role, not a tenant-store operator. It must not be used to bypass store-scoped controls. Inventory is a back-office role for inventory and inventory reporting; it is redirected away from POS and admin-only modules.")
    add_heading(doc, "2.2 Environment Permission Matrix", 2)
    env_rows = [
        ("Cashier", "Execute POS in assigned store", "Test user only when invited", "None", "None", "None"),
        ("Inventory", "Edit inventory; inventory reports", "Test user only", "None", "None", "None"),
        ("Admin / System Owner", "Full tenant operations", "Approve business acceptance", "Read approved releases", "None", "Request restore; no direct vault access"),
        ("Super Admin", "Invite governance only", "Validate invite workflow", "None", "None", "None"),
        ("System Administrator", "Operational support; temporary audited elevation", "Full technical control", "Edit/deploy via approved change", "Manage functions/secrets", "Restore operator"),
        ("IT Manager", "Read/approve emergency access", "Approve release", "Approve merge/release", "Approve privileged changes", "Approve restore and retention"),
        ("Developer", "No routine access; time-bound support only", "Edit", "Edit via pull request", "Deploy only if delegated", "None"),
        ("Privacy Contact / DPO", "Read only when investigating a privacy case", "None", "None", "None", "Approve disclosure/deletion handling"),
    ]
    add_table(doc, ["Role", "Production App/Data", "Staging/Test", "Source & CI", "Edge Functions", "Backups"], env_rows,
              [1.3, 1.45, 1.05, 1.0, 0.9, 0.8], font_size=7.2)
    add_heading(doc, "2.3 Roles, Responsibilities, and Accountability", 2)
    role_rows = [
        ("Admin / System Owner", "Store configuration, operational data quality, staff access, daily control review", "User creation/removal, stock corrections, purchase orders, routine settings", "Accountable for store impact, unauthorized access caused by poor provisioning, and business acceptance"),
        ("Cashier", "Accurate cart, payment, receipt, and end-of-shift reporting", "Customer confirms tender; Admin approves void/refund exceptions", "Accountable for transactions under own credentials and immediate reporting of discrepancies"),
        ("Inventory Staff", "Product master data, counts, stock adjustments, inventory reports", "Admin approves unusual or high-impact adjustments", "Accountable for count accuracy, reason codes, and supporting notes"),
        ("Super Admin", "Invite-code governance and platform onboarding controls", "Approves/generates governed invite access", "Accountable for invite misuse or excessive issuance within assigned authority"),
        ("Store Team Lead", "First-line incident intake, verification, workaround, and user communication", "Escalates unresolved or multi-user problems", "Accountable for complete ticket details and timely escalation"),
        ("System Administrator", "Availability, Supabase configuration, deployments, logs, restores, technical diagnosis", "Implements approved normal changes; emergency action per incident authority", "Accountable for technical recovery, change evidence, and privileged access records"),
        ("IT Manager", "Service ownership, risk acceptance, major-incident command, vendor coordination", "Approves releases, emergency change, restore, and service communications", "Ultimately accountable for service continuity and post-incident action"),
        ("Privacy Contact / DPO", "Rights requests, breach assessment, privacy advice, records of processing", "Approves privacy response and disclosure decisions", "Accountable for lawful, transparent, and timely privacy handling"),
    ]
    add_table(doc, ["Role", "Owns", "Approves", "Accountable when something goes wrong"], role_rows,
              [1.25, 1.75, 1.5, 2.0], font_size=7.6)
    add_heading(doc, "2.4 Control Rules", 2)
    add_bullets(doc, [
        "No shared accounts. Every action must be attributable to one named user.",
        "Access is granted by role and store membership, reviewed quarterly, and removed immediately on separation or role change.",
        "Privileged technical access is time-bound, ticketed, approved, and reviewed after use.",
        "The person implementing a high-risk change should not be the sole approver of that change.",
        "Audit logs are append-only evidence; corrections are recorded as new events rather than silent edits.",
    ])
    add_heading(doc, "2.5 Quarterly Access Review Procedure", 2)
    add_steps(doc, [
        "Admin exports or reviews the current user list, roles, active status, and store assignments.",
        "Each Store Team Lead confirms which users still require access and whether duties changed.",
        "Admin removes leavers, corrects excessive store scope, and records the approval/effective date.",
        "System Administrator reviews privileged technical accounts, tokens, secrets, and emergency access evidence.",
        "IT Manager signs off unresolved exceptions and assigns a due date; the review record is retained as audit evidence.",
    ])
    add_heading(doc, "Approval Triggers Outside the Review Cycle", 3)
    add_body(doc, "Immediate review is required after termination, transfer, suspected compromise, major incident, new store assignment, introduction of a new role, or a material RLS/route change. Emergency access expires when the incident ends and must not silently become permanent.")

    # 3 ITIL escalation
    new_section(doc, "3. ITIL Artifact — Escalation Path", "A practical service desk route tied to the roles and accountability model")
    add_heading(doc, "3.1 Escalation Diagram", 2)
    escalation_steps = [
        ("Tier 0 — User self-check", "Cashier / Inventory / Admin checks store selection, connectivity, device power, credentials, and the troubleshooting steps in Section 4. Channel: in-app message and this manual."),
        ("Tier 1 — Store Team Lead / Admin", "Confirms impact, protects cash and records, gathers screenshots/time/store/user/transaction ID, applies approved workaround, and opens a support ticket. Channel: ticketing system; urgent call or group chat."),
        ("Tier 2 — System Administrator", "Investigates logs, RLS, database, Edge Functions, offline queue, deployment, browser/hardware integration, and data correction options. Channel: ticket + technical incident channel."),
        ("Tier 3 — IT Manager / Privacy Contact / Vendor", "Commands major incidents, approves emergency change or restore, coordinates Supabase/hosting vendor, and leads privacy assessment when personal data may be affected. Channel: emergency bridge, phone, email, vendor portal."),
    ]
    add_flow_diagram(doc, escalation_steps)
    add_heading(doc, "3.2 Tier Criteria and Targets", 2)
    tier_rows = [
        ("Tier 0", "Single-user usability issue; known warning; no data loss or security concern", "Immediate / up to 10 min", "Manual, in-app feedback", "User"),
        ("Tier 1", "Issue persists after basic checks; cash discrepancy; one terminal/store affected; access request", "Acknowledge 15 min (P1/P2), 1 business hr (P3); target workaround 1 hr", "Ticket; urgent chat/call", "Store Team Lead / Admin"),
        ("Tier 2", "Multiple users; repeated failed sync; suspected RLS/database defect; checkout unavailable; incorrect stock/sale state", "Acknowledge 30 min; P1 restore/workaround target 4 hrs; P2 1 business day", "Ticket + technical incident channel", "System Administrator"),
        ("Tier 3", "Major outage, restore/emergency change, vendor dependency, confirmed/suspected breach, legal or reputational impact", "P1 acknowledge 15 min; executive update every 60 min; recovery target set by incident commander", "Emergency bridge/phone/email/vendor portal", "IT Manager; Privacy Contact for privacy"),
    ]
    add_table(doc, ["Tier", "Escalate when", "Response / resolution target", "Channels", "Role"], tier_rows,
              [0.55, 2.0, 1.55, 1.2, 1.2], font_size=7.4)
    add_body(doc, "Targets are operational commitments for the capstone deployment and should be replaced by an approved service-level agreement if Kodigo is commercialized. Resolution depends on cause; when the target cannot be met, the owner must provide an updated estimate and workaround.")
    add_heading(doc, "3.3 Severity Guide", 2)
    severity_rows = [
        ("P1 Critical", "Checkout unavailable for a store/all stores; active security incident; material data corruption", "Immediate Tier 2 and Tier 3; preserve logs; activate continuity process"),
        ("P2 High", "Major feature unavailable, repeated sync failures, wrong totals/stock affecting operations", "Tier 1 validates, then Tier 2 within 15 minutes"),
        ("P3 Normal", "Single-user problem with workaround; report/export or noncritical device issue", "Tier 1 ticket; escalate if unresolved within 1 business day"),
        ("P4 Request", "How-to question, enhancement, planned access or configuration change", "Normal service request queue and change process"),
    ]
    add_table(doc, ["Priority", "Example", "Action"], severity_rows, [1.0, 3.0, 2.5], font_size=8.5)
    add_heading(doc, "3.4 Required Ticket Information", 2)
    add_bullets(doc, [
        "Reporter name and role; store and terminal; date/time and timezone.",
        "What the user attempted, expected result, actual result, and exact message.",
        "Transaction/receipt/product ID when applicable; never paste passwords or full payment credentials.",
        "Screenshot with unnecessary personal data redacted; browser/device and online/offline status.",
        "Business impact, number of users affected, workaround tried, and whether data or privacy may be at risk.",
    ])

    # 4 Full manual
    new_section(doc, "4. Full Manual Guide", "Complete operating instructions for Kodigo’s core workflows")
    add_heading(doc, "4.1 Manual Rationale and Training Approach", 2)
    add_body(doc, "This manual is for sari-sari store owners/administrators, cashiers, inventory staff, super administrators, and first-line support personnel using Kodigo across one or more branches. Their tasks differ, but they share the same operational data: a cashier’s sale changes inventory; inventory changes affect checkout; and administrators use both to replenish and report.")
    add_body(doc, "A Full Manual is the right primary documentation type because the system combines routine high-frequency work with lower-frequency, higher-risk tasks. A quick-start guide would help with first login or checkout but would omit role boundaries, offline behavior, stock correction, reporting, escalation, and policy obligations. A video is harder to search during a live transaction and becomes stale whenever screens change. An interactive simulation is useful for practice but expensive to maintain and cannot serve as the authoritative control and troubleshooting reference. This manual remains searchable, printable, reviewable, and suitable for training and audit evidence.")
    add_body(doc, "Training should use a blended model: instructor-led onboarding for Admins and new Cashiers, a self-paced copy of this manual for reinforcement, and just-in-time checklists at the terminal. Instructor-led practice is appropriate for payment, void/refund, stock adjustment, and incident scenarios because mistakes affect cash and records. Self-paced review supports shift-based staff, while just-in-time steps reduce memory burden during real work. Competency should be demonstrated in a staging store before production access is granted.")
    add_callout(doc, "Recommended training evidence", "Record attendance, role assigned, practice scenarios completed, trainer sign-off, and date of refresher training. Re-train after major workflow or policy changes.")

    add_heading(doc, "4.2 What Kodigo Does", 2)
    add_body(doc, "Kodigo is a cloud-first, multi-store retail operations system. It supports checkout and receipts, inventory and stock adjustments, supplier and purchase-order management, stock alerts, dashboards, analytics, rankings, reporting, store branding, user management, notifications, and account security. Supabase provides authentication and store-scoped database access; IndexedDB supports selected offline queues; Web Serial can trigger a compatible cash drawer.")
    add_heading(doc, "4.3 Before You Begin", 2)
    add_bullets(doc, [
        "Use a supported modern browser. Cash-drawer Web Serial requires a compatible browser and HTTPS or localhost.",
        "Confirm the device clock, network, receipt printer/scanner connection, and correct store assignment.",
        "Use your own account. Do not share a password or leave an unlocked terminal unattended.",
        "For POS, select one specific active store. The All Stores view is for consolidated review only.",
        "Know your role: Cashier uses POS; Inventory manages products/stock and inventory reports; Admin manages store operations; Super Admin governs invites.",
    ])
    add_heading(doc, "4.4 Sign In, Store Context, and Account Security", 2)
    add_steps(doc, [
        "Open the Kodigo login page and enter your assigned email and password.",
        "After login, verify the name/branding and active store shown in the interface. If you are assigned to several stores, choose the branch you are operating.",
        "If you see no store or the wrong store, stop before posting data and ask the Admin to check membership.",
        "To change your password, open Account Security, enter the current password, then a new password of at least eight characters and confirm it.",
        "When finished, use Logout. Closing the browser alone is not a substitute on a shared device.",
    ])
    add_callout(doc, "Password reset", "Use Forgot Password from the login page or ask an Admin to send a managed-user reset. Support must never ask for the current password.", fill=AMBER)

    add_heading(doc, "4.5 POS Checkout", 2)
    add_steps(doc, [
        "Open POS Terminal and confirm the active store. Admin POS access is intended for desktop; Cashier is the frontline POS role.",
        "Find an item by product name, SKU, or barcode. F2 focuses search. A compatible keyboard-emulating scanner can input the barcode.",
        "Add the product to the cart. Adjust quantity only after physically confirming the requested amount and available stock. Numeric pre-entry can set a quantity multiplier.",
        "Review product, quantity, unit price, subtotal, discount/tax presentation, and total. Remove accidental items before payment.",
        "Select Charge or press F9 when the cart contains items. Choose the payment method and enter cash received where required.",
        "Confirm the change due with the customer, then complete checkout once. Wait for confirmation—do not double-click or repeat a sale because the screen is slow.",
        "Issue the generated receipt. If configured, print it and open the cash drawer through the approved hardware action.",
        "If the sale is queued offline, note the terminal and time. Keep the browser data intact until synchronization succeeds.",
    ])
    add_heading(doc, "Payment and Receipt Checks", 3)
    add_bullets(doc, [
        "Cash: count tender aloud, confirm amount entered, and return the displayed change.",
        "GCash/card/bank transfer/other: record only the supported method/reference; do not store full card numbers, PINs, OTPs, or wallet passwords.",
        "Receipt: verify store identity, invoice/receipt number, items, totals, payment, cashier, date/time, and VAT/non-VAT information configured for the store.",
    ])
    add_heading(doc, "Voids, Refunds, Returns, and Closeout", 3)
    add_body(doc, "These actions affect revenue, stock, cash, and audit evidence. Follow the organization’s approval threshold. The Cashier records the reason and requests Admin/Team Lead approval; the Admin verifies the original transaction and supporting facts. Never delete a transaction to hide an error. Kodigo’s transaction lifecycle records status/events so corrections remain traceable.")

    add_heading(doc, "4.6 Offline Operation and Synchronization", 2)
    add_body(doc, "Kodigo caches products for offline POS lookup and can queue eligible offline sales and selected mutations in the browser’s IndexedDB. When connection returns, focus/visibility/online events trigger synchronization and data revalidation.")
    add_steps(doc, [
        "If an offline message appears, confirm the correct store and continue only if the screen explicitly allows the operation.",
        "Do not clear browser data, switch browser profiles, uninstall the browser, or use private mode while transactions are pending.",
        "Reconnect to the approved network and leave Kodigo open. Allow the queue to replay and watch for success or rejection messages.",
        "If an RLS/database error appears, do not repeatedly retry. Capture the message and escalate because business-rule errors are not safely queued as network failures.",
        "Reconcile pending receipts and stock after synchronization. Escalate any duplicate, missing, or rejected sale immediately.",
    ])

    add_heading(doc, "4.7 Inventory and Product Management", 2)
    add_steps(doc, [
        "Open Inventory and select a specific store. Admin and Inventory roles can access this area.",
        "To add a product, enter a unique SKU, optional barcode, product name, category, selling unit, cost, price, opening stock, low-stock threshold, reorder level, lead time, and supplier where applicable.",
        "Review cost and selling price before saving. A wrong unit or conversion factor can distort stock and margin reports.",
        "To edit, open the product, change only verified fields, and save. Avoid changing SKU/barcode identity during active checkout without coordination.",
        "For stock changes, use Stock Adjustment rather than overwriting history. Enter the delta, reason, and a meaningful note; verify before/after quantities.",
        "Delete/deactivate products only after checking links and operational need. Use the confirmation dialog and preserve history required for reports.",
    ])
    add_callout(doc, "Bulk/per-piece selling", "Stock is intended to be tracked in the smallest sellable unit. Validate purchase-unit conversion data before relying on it for production reports because optional conversion display fields may require deployment-specific verification.", fill=AMBER)

    add_heading(doc, "4.8 Stock Alerts and Restocking", 2)
    add_steps(doc, [
        "Review low, critical, and out-of-stock alerts from Notifications/Restocking.",
        "Confirm physical count before ordering. The suggested quantity is a planning aid, not an automatic approval.",
        "Review the proposed quantity, reorder level, lead time, recent sales, and supplier availability.",
        "Create purchase orders grouped by supplier/store. Validate item, quantity, unit cost, supplier, and expected delivery.",
        "Move the order through Sent, Received, or Cancelled status. Mark Received only after checking delivered goods and quantities.",
        "Investigate supplier-score changes and delivery exceptions; add supporting records rather than changing history without explanation.",
    ])
    add_body(doc, "The current planning rule is: suggested quantity = max(reorder level × 2 − current stock, reorder level). Human review remains required because promotions, seasonality, shelf capacity, cash limits, and supplier constraints are not fully captured by the formula.")

    add_heading(doc, "4.9 Suppliers", 2)
    add_steps(doc, [
        "Open Suppliers and select the relevant store context.",
        "Add or edit supplier name, contact person, phone/email/address, lead time, and related product information.",
        "Use purchase-order and delivery history to evaluate reliability, price, and overall score.",
        "Before deletion, check linked purchase orders. Retain records needed for financial, operational, or audit purposes.",
    ])

    add_heading(doc, "4.10 Dashboard, Analytics, Rankings, and Reports", 2)
    add_bullets(doc, [
        "Dashboard: reviews daily revenue, transaction count, average order value, estimated profit, recent transactions, and stock conditions.",
        "Analytics: filters Today, 7, 30, or 90 days and analyzes revenue/profit, hourly sales, category sales, and recent transactions.",
        "Rankings: aggregates sale items to rank products by sales activity. Use this instead of the dashboard Top Products panel when a sales-based ranking is required.",
        "Reports: Admin can generate sales and inventory workbooks with date, payment, status, cashier, product, category, and selling-unit filters. Inventory can export inventory/stock-focused reports only.",
    ])
    add_steps(doc, [
        "Select one store or an authorized All Stores aggregate where the report supports it.",
        "Choose the date range and filters; refresh and review summary totals before export.",
        "Confirm status filters include/exclude voided or refunded transactions as intended.",
        "Export only to an approved location. Apply the data-classification rules in Section 6 and securely delete obsolete copies.",
    ])

    add_heading(doc, "4.11 Store Settings, Branding, Users, and Notifications", 2)
    add_bullets(doc, [
        "General/branding: Admin maintains store name, address, tax rate, registered/business name, TIN, branch code, VAT status, receipt label/prefix, terminal identifier, permits/accreditation text, contact details, and logo.",
        "Users: Admin can create, edit, reset, or remove Admin/Cashier/Inventory users only within shared assigned stores. Multi-store Admin assignments must be deliberate.",
        "Notifications: Admin can manage preferences for low stock, out of stock, daily summary, and sales milestones where configured.",
        "Super Admin: governs invite codes and system onboarding, not tenant operational CRUD.",
    ])
    add_callout(doc, "User removal", "Before removing or changing a user, verify identity, approval, active shifts, pending offline work, and store assignments. Access loss should be immediate, but operational records remain attributable to the former user.")

    add_heading(doc, "4.12 Basic Troubleshooting", 2)
    trouble_rows = [
        ("Cannot sign in", "Wrong password, reset required, missing profile/role", "Retry carefully; use Forgot Password; ask Admin to confirm the account. Escalate if role cannot be resolved."),
        ("No store / wrong store", "Missing membership or stale selection", "Stop posting; refresh; re-login; ask Admin to verify store assignment."),
        ("POS unavailable", "Role/device restriction, network, deployment", "Cashier/Admin desktop only as allowed; check network; capture error; use approved continuity process."),
        ("Barcode not found", "Wrong barcode, product inactive/missing, scanner mode", "Search by name/SKU; verify product and scanner input; do not create duplicates."),
        ("Sale appears pending", "Offline queue or slow connection", "Do not repeat sale; reconnect and wait for sync; reconcile by receipt/time before retry."),
        ("Stock is incorrect", "Unreceived PO, unsynced sale, wrong adjustment", "Count physically; review sale/adjustment/PO history; Admin approves a documented correction."),
        ("Report totals differ", "Store/date/status filter or pending sync", "Confirm context and filters; sync; compare transaction statuses; escalate with export and IDs."),
        ("Cash drawer fails", "Unsupported browser, insecure context, permission/cable", "Use HTTPS/localhost and compatible browser; reconnect; allow permission after user gesture; use manual fallback."),
        ("Access denied/RLS error", "Role/store policy blocks action", "Do not bypass or keep retrying; verify assignment; send exact error to Tier 2."),
        ("App looks stale", "Cached data or inactive tab", "Bring tab to front, refresh, check online state, and re-login if safe; preserve pending offline data."),
    ]
    add_table(doc, ["Problem", "Likely cause", "Safe action"], trouble_rows, [1.35, 1.9, 3.25], font_size=8.0)
    add_heading(doc, "4.13 End-of-Shift Checklist", 2)
    add_bullets(doc, [
        "Complete or cancel open carts; do not leave an ambiguous payment screen.",
        "Confirm pending offline transactions are synchronized or formally handed over.",
        "Perform the approved cashier closeout and explain any variance.",
        "Secure receipts/reports and remove customer-visible information from the counter.",
        "Log out, lock the device, and report unresolved incidents with ticket number to the next shift.",
    ])

    # 5 policies
    new_section(doc, "5. Acceptable Use and BYOD Policy", "Adoptable policy text for Kodigo users and devices")
    add_heading(doc, "5.1 Acceptable Use Policy (AUP)", 2)
    add_heading(doc, "Purpose", 3)
    add_body(doc, "This policy protects Kodigo, the stores it supports, and the people whose information is processed. It establishes authorized, secure, and accountable use of the application, accounts, devices, network access, reports, and connected peripherals.")
    add_heading(doc, "Scope", 3)
    add_body(doc, "This policy applies to all employees, owners, contractors, students/testers, support personnel, and third parties who access Kodigo or its data in production, staging, test, backup, export, or support environments.")
    add_heading(doc, "Allowed Uses", 3)
    add_bullets(doc, [
        "Perform assigned retail, inventory, reporting, administration, support, or audit duties within the user’s approved role and stores.",
        "Use authorized scanners, receipt printers, cash drawers, browsers, and networks for business operations.",
        "Create exports only for an approved business purpose and store them in an approved restricted location.",
        "Report errors, vulnerabilities, lost devices, suspicious access, and privacy concerns promptly through the escalation path.",
        "Use staging/test accounts and non-production data for training and testing whenever feasible.",
    ])
    add_heading(doc, "Prohibited Uses", 3)
    add_bullets(doc, [
        "Sharing accounts, passwords, reset links, OTPs, invite codes, service-role keys, or sessions; using another person’s account; or leaving a terminal unlocked.",
        "Accessing another store, role, user record, report, or log without authorization, even if a technical path appears available.",
        "Changing, deleting, fabricating, or concealing sales, stock, supplier, user, receipt, closeout, error, or audit records to misrepresent events.",
        "Installing unapproved extensions/software; bypassing security controls; probing production; or connecting unknown USB/serial devices.",
        "Saving full card numbers, PINs, OTPs, wallet passwords, government IDs, or unrelated customer information in notes, logs, screenshots, or tickets.",
        "Uploading Kodigo data to personal email, public cloud storage, social media, public AI tools, or removable media without written approval.",
        "Clearing browser data when offline transactions are pending, or intentionally disrupting synchronization, backups, logging, or availability controls.",
    ])
    add_heading(doc, "Consequences", 3)
    add_body(doc, "Violations may result in immediate access suspension, preservation and review of relevant logs, corrective training, disciplinary action, contract consequences, recovery of losses where lawful, and referral to management, the Privacy Contact, law enforcement, or regulators as appropriate. Actions will be proportionate, documented, and consistent with applicable law and organizational rules.")
    add_callout(doc, "CIA Triad linkage", "The rule against credential sharing protects Confidentiality by limiting data to authorized users, protects Integrity by preserving reliable attribution, and supports Availability by reducing account compromise and service disruption.", fill=GREEN)

    add_heading(doc, "5.2 Bring Your Own Device (BYOD) Policy", 2)
    add_body(doc, "BYOD is applicable only when the organization expressly permits personal devices for Kodigo. A personal device is not automatically authorized merely because the web application can open on it.")
    add_heading(doc, "Device Security Requirements", 3)
    add_bullets(doc, [
        "Supported, legally licensed operating system and browser with automatic security updates enabled.",
        "Device passcode of at least six digits or strong password; biometric unlock may supplement but not replace the passcode.",
        "Automatic screen lock within five minutes; full-device encryption; active anti-malware where supported; firewall enabled.",
        "No rooted/jailbroken devices, unsupported operating systems, shared family profiles, or unapproved browser extensions.",
        "Use only trusted networks or an organization-approved VPN; public/shared Wi-Fi must not be used for administrative or export activity without protection.",
        "Do not locally save reports, receipts, screenshots, or credentials unless explicitly approved. Pending offline data must be synchronized promptly.",
        "Lost/stolen devices, suspected compromise, or unauthorized family/third-party access must be reported immediately and no later than one hour after discovery.",
    ])
    add_heading(doc, "Support Boundaries", 3)
    add_body(doc, "Support covers Kodigo configuration, supported browsers, approved scanner/printer/cash-drawer integration, and reasonable diagnosis. The organization does not guarantee support for personal hardware defects, data plans, home routers, unrelated apps, personal backups, or unsupported operating systems. Support may require the user to reproduce the issue on an organization-managed device.")
    add_heading(doc, "Liability and Privacy", 3)
    add_body(doc, "The device owner remains responsible for personal-device costs, lawful content, physical care, and personal backups. The organization is responsible for protecting Kodigo business data under its control but is not liable for ordinary loss of personal data caused by user failure to back it up, except where law or contract provides otherwise. Investigation and remote session revocation may affect Kodigo data or sessions; personal content will not be inspected beyond what is necessary and authorized for security, legal, or privacy response.")
    add_heading(doc, "Offboarding and Device Disposal", 3)
    add_steps(doc, [
        "Admin removes or changes Kodigo access and store assignments on the effective separation/role-change date.",
        "User logs out, removes saved credentials/bookmarks that reveal restricted endpoints, and synchronizes or hands over pending authorized work.",
        "System Administrator revokes active sessions/tokens where available and confirms no pending offline queue remains on the device.",
        "User securely deletes approved local exports, receipts, caches, and screenshots; Admin verifies completion through an offboarding checklist.",
        "Before selling, recycling, or transferring the device, the owner performs a secure factory reset after preserving lawful personal backups.",
    ])
    add_heading(doc, "Exceptions and Acknowledgment", 3)
    add_body(doc, "Any exception must identify the business need, data involved, compensating controls, owner, approval, and expiry date. The IT Manager approves security exceptions; the Privacy Contact also approves exceptions involving personal data. Unapproved exceptions are policy violations.")
    add_table(doc, ["User acknowledgment", "Record"], [
        ("I understand that Kodigo access is limited to my role and assigned stores.", "Name / signature / date"),
        ("I will report loss, compromise, incorrect records, or pending offline work promptly.", "Training record or policy acceptance"),
        ("I consent to proportionate security controls on authorized BYOD access and will complete offboarding.", "Device ID and approval expiry"),
    ], [4.6, 1.9], font_size=8.7)

    # 6 data classification
    new_section(doc, "6. Data Classification Table", "Inventory of Kodigo data and minimum protection expectations")
    add_body(doc, "Classification uses three practical labels: Personal Data (PD), Sensitive Personal Information (SPI), and Operational/Non-personal (OP). Credentials and security telemetry may not always fall within a statutory SPI category, but Kodigo treats them as Restricted because compromise creates significant risk. Customer identity is not required for ordinary cash sales; do not collect it unless a defined feature and lawful purpose require it.")
    data_rows = [
        ("User name, profile ID, avatar", "PD", "Confidential", "Store-scoped access; authenticated use; correct/erase through Admin/privacy process"),
        ("User email and phone (if used)", "PD", "Confidential", "Use for account/contact only; mask in tickets and exports; encrypted transit/storage"),
        ("Password/authentication tokens", "PD (security credential)", "Restricted", "Supabase Auth only; never public tables/logs; no sharing; reset and revoke on compromise"),
        ("Role and store memberships", "PD", "Confidential", "Least privilege; Admin shared-store scope; quarterly review and immediate offboarding"),
        ("IP address, browser/device, error context", "PD when linkable", "Confidential", "Collect only for security/support; restrict logs; time-limited retention"),
        ("Government identifiers if later collected", "SPI", "Highly Restricted", "Do not collect by default; documented legal basis, encryption, narrow access, retention limit"),
        ("Health/biometric information", "SPI", "Highly Restricted", "Not required by Kodigo; prohibited unless separately approved and lawfully implemented"),
        ("Store TIN, permit/accreditation details", "OP; may become PD for sole proprietor", "Confidential", "Admin-only edit; controlled receipt display; accurate and current"),
        ("Store name, address, contact, branding", "OP / possible PD", "Internal–Confidential", "Admin control; publish only intended receipt/contact fields"),
        ("Products, SKU/barcode, cost, price, stock", "OP", "Confidential", "Store-scoped RLS; change history/adjustments; backups"),
        ("Supplier contact person, email, phone, address", "PD + OP", "Confidential", "Business-purpose use; store scope; limit exports and deletion to retention rules"),
        ("Sales, sale items, totals, discounts, tax", "OP; PD if linked to person", "Confidential", "Store scope; append-only lifecycle evidence; financial retention; encrypted transfer"),
        ("Payment method/reference", "PD/financial context when linkable", "Restricted", "Never store full card/PIN/OTP; mask reference; narrow access and retention"),
        ("Receipts/invoice numbers", "OP; may contain PD", "Confidential", "Issue to customer; protect duplicate copies/exports; approved retention"),
        ("Stock adjustments and creator label", "PD + OP", "Confidential", "Named attribution, reason/note, immutable audit trail, store scope"),
        ("Purchase orders and supplier scores", "OP; may contain PD contacts", "Confidential", "Admin scope; accuracy review; protected exports"),
        ("Audit logs, error logs, sale events, closeouts", "PD + OP", "Restricted", "Append-only; Admin/technical need-to-know; monitor access; defined retention"),
        ("Backups and exported reports", "Same as source data", "Restricted", "Encrypted/restricted vault; restore testing; minimum 30-day backup retention per runbook; controlled disposal"),
        ("Invite codes and reset links", "PD/security secret", "Restricted", "Single-purpose, limited validity/use, never public or reused, revoke when exposed"),
    ]
    add_table(doc, ["Data item", "Classification", "Protection", "Minimum controls"], data_rows,
              [1.65, 1.15, 1.0, 2.7], font_size=7.25)
    add_heading(doc, "Handling Rules by Protection Level", 2)
    add_table(doc, ["Level", "Handling rule"], [
        ("Internal", "Authenticated business use; share only within the organization for a legitimate task."),
        ("Confidential", "Role/store-scoped access, encrypted transfer, approved storage/export, logged changes, secure disposal."),
        ("Restricted / Highly Restricted", "Named need-to-know access, no casual export, additional approval, monitoring, short retention where feasible, immediate incident response."),
    ], [1.5, 5.0], font_size=9)
    add_heading(doc, "Retention and Disposal Baseline", 2)
    add_table(doc, ["Record group", "Baseline decision"], [
        ("Sales, receipts, tax and payment records", "Retain for the approved financial/legal period; lock against routine deletion; dispose securely when authorized."),
        ("User/account and membership records", "Retain while active plus the period needed for audit or claims; revoke access immediately on offboarding."),
        ("Audit, error and security logs", "Keep long enough to investigate incidents and demonstrate control operation; restrict and purge on an approved schedule."),
        ("Exports and screenshots", "Keep only for the specific task; delete from downloads, email, chat, and removable media when the task/retention need ends."),
        ("Backups", "Minimum 30-day restricted retention per current runbook, subject to the production backup plan and legal requirements."),
    ], [2.15, 4.35], font_size=8.3)

    # 7 privacy
    new_section(doc, "7. Privacy Notice", "Short notice for Kodigo account holders and business contacts")
    add_callout(doc, "Plain-language notice", "This notice should be displayed at account creation or made accessible from the login/settings area. The deploying organization is the personal information controller and must insert its legal name and any legally required registration details before publication.")
    privacy_points = [
        ("1. What we collect", "We collect account details such as your name, email, role and assigned store; activity records such as sales, stock adjustments, reports and audit events; and limited device/error information needed to secure and support Kodigo. Ordinary cash sales do not require customer identity."),
        ("2. Why we use it", "We use the information to authenticate users, operate store transactions and inventory, issue receipts, manage staff access, produce reports, prevent misuse, troubleshoot errors, meet recordkeeping duties, and improve service reliability. We will not use it for an unrelated purpose without notice and an appropriate lawful basis."),
        ("3. How we protect and share it", "Access is limited by role and assigned store, protected by authentication, encrypted connections, row-level security, logs, backups, and controlled support access. We share data only with authorized personnel and service providers needed to run or support Kodigo, subject to appropriate safeguards and legal requirements."),
        ("4. Your choices and rights", "Subject to applicable law, you may ask to access or correct your data, object to or restrict certain processing, request deletion when retention is no longer required, withdraw consent where processing relies on consent, and raise a complaint. We may verify identity and retain records required for security, transactions, or law."),
        ("5. Contact us", "Send privacy questions or requests to the Store Administrator / Privacy Contact using the official store email shown in Kodigo Settings and on the store receipt. If unavailable, submit a support ticket marked “Privacy Request.” Do not include passwords, OTPs, or unnecessary identity documents in the first message."),
    ]
    for title, text in privacy_points:
        add_heading(doc, title, 2)
        add_body(doc, text)
    add_heading(doc, "Privacy Principles Applied", 2)
    add_table(doc, ["Principle", "How it is applied"], [
        ("Collection limitation", "Kodigo collects only data needed for accounts, store operations, security, support, and required records; customer identity and SPI are not required by default."),
        ("Purpose specification", "The notice states the operational, access-control, reporting, security, support, and recordkeeping purposes at collection."),
        ("Use limitation", "Data is not used or disclosed for unrelated purposes without notice, authority, and safeguards; access is role/store-scoped."),
        ("Transparency", "Users are told what is collected, why, how it is protected/shared, their rights, and how to contact the responsible organization."),
    ], [1.55, 4.95], font_size=9)
    add_heading(doc, "Privacy Request Handling", 2)
    add_steps(doc, [
        "Log the request and acknowledge it through an approved channel without asking for unnecessary identity documents.",
        "Verify the requester proportionately, identify the relevant stores/systems, and preserve required transaction or security records.",
        "The Privacy Contact decides the lawful response, coordinates Admin/System Administrator action, and records any limitation or refusal reason.",
        "Provide the response securely and record completion, disclosure recipients, corrections, deletions, or follow-up commitments.",
    ])
    add_callout(doc, "Privacy incident", "If data may have been accessed, changed, lost, or disclosed without authority, preserve evidence and escalate directly to Tier 3. Do not notify affected people or external parties without the Privacy Contact/incident commander’s coordinated decision.", fill=AMBER)

    # Appendix
    new_section(doc, "Appendix A. Control and Evidence Notes", "Technical anchors used to tailor this document to the current Kodigo implementation")
    evidence_rows = [
        ("Application and routes", "kodigo-ui/src/App.tsx; Sidebar.tsx; types/index.ts", "Admin, Cashier, Inventory, Super Admin; route guards and module visibility"),
        ("Authentication and tenancy", "kodigo-ui/src/stores/authStore.ts; migration chain", "Supabase Auth, profiles, store_users, active store, row-level security"),
        ("POS and offline", "cartStore.ts; transactions.ts; offline-sync.ts; hardware.ts", "Checkout, queued sales/mutations, network-vs-database errors, cash drawer"),
        ("Inventory/suppliers/reporting", "productStore.ts; supplierStore.ts; reporting.ts; relevant pages", "Product/stock, supplier/PO, reports and exports"),
        ("Production controls", "docs/production-ops.md; migrations 18–28", "Backups, restore, deployment, rollback, audit/error logs, user RBAC hardening"),
    ]
    add_table(doc, ["Area", "Repository evidence", "Document implication"], evidence_rows, [1.35, 2.5, 2.65], font_size=8)
    add_heading(doc, "Deployment Decisions Still Required", 2)
    add_bullets(doc, [
        "Insert the deploying organization’s legal identity and official Privacy Contact details before publishing the privacy notice.",
        "Approve concrete service hours, incident targets, access-review frequency, export retention, and legal record-retention schedule.",
        "Confirm the supported invite-generation path, production environment ownership, backup subscription/retention, and vendor escalation contacts.",
        "Validate selling-unit conversion reporting and every policy/role against the exact production migration state before go-live.",
    ])
    add_heading(doc, "Approval Record", 2)
    add_table(doc, ["Role", "Name", "Signature", "Date"], [
        ("System Owner", "", "", ""),
        ("IT Manager / Technical Adviser", "", "", ""),
        ("Privacy Contact", "", "", ""),
        ("Capstone Adviser", "", "", ""),
    ], [1.8, 1.8, 1.8, 1.1], font_size=9)

    # Reapply header/footer to all sections and save.
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
        section.different_first_page_header_footer = True
        add_header_footer(section)

    doc.settings.odd_and_even_pages_header_footer = False
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
